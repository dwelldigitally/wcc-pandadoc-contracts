"""
WCC Master Student Enrollment Contract Template
Generates a DOCX file formatted for PandaDoc template upload.

Token syntax: [TokenName] — PandaDoc replaces these with CRM/API values
Field tags: These are added in PandaDoc's editor after upload (signature, date, checkbox, dropdown)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)

# ── Helper Functions ──
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(139, 0, 0)  # Dark red (WCC brand)
    return h

def add_field_row(table, label, token_or_value, bold_label=True):
    row = table.add_row()
    cell_label = row.cells[0]
    cell_value = row.cells[1]
    p_label = cell_label.paragraphs[0]
    run = p_label.add_run(label)
    run.bold = bold_label
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    p_value = cell_value.paragraphs[0]
    run_v = p_value.add_run(token_or_value)
    run_v.font.size = Pt(10)
    run_v.font.name = 'Arial'
    return row

def add_simple_table(headers_data):
    """Add a simple 2-column table with label:value pairs"""
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Set column widths
    for row_data in headers_data:
        add_field_row(table, row_data[0], row_data[1])
    return table

def add_static_paragraph(text, bold=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Arial'
    return p

def add_spacer():
    doc.add_paragraph()


# ════════════════════════════════════════════════════════════
# PAGE 1: HEADER & INSTITUTION INFO
# ════════════════════════════════════════════════════════════

# Logo placeholder
p_logo = doc.add_paragraph()
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_logo.add_run('[WCC LOGO - Upload in PandaDoc Editor]')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(128, 128, 128)
run.italic = True

# Institution header
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Western Community College Inc.')
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('WESTERN COMMUNITY COLLEGE')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor(139, 0, 0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('#201 8313 120th Street, Surrey, BC V3W 3N4')
run.font.size = Pt(10)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('STUDENT ENROLLMENT CONTRACT')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(139, 0, 0)

add_spacer()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Western Community College is DESIGNATED by the Private Training Institutions Regulatory Unit (PTIRU)')
run.font.size = Pt(9)
run.italic = True

add_spacer()

# ════════════════════════════════════════════════════════════
# SECTION 1: STUDENT INFORMATION
# ════════════════════════════════════════════════════════════

add_heading_styled('STUDENT INFORMATION', level=2)

student_fields = [
    ('Last Name:', '[Student.LastName]'),
    ('First Name & Middle Name:', '[Student.FirstName]'),
    ('Usual First Name:', '[Student.UsualFirstName]'),
    ('Student Email Address:', '[Student.Email]'),
    ('Student Telephone Number:', '[Student.Phone]'),
    ('Personal Education Number (if available):', '[PEN — Advisor enters in PandaDoc]'),
    ('Mailing Address:', '[Student.MailingAddress]'),
    ('Date of Birth:', '[Student.DateOfBirth]'),
    ('Mailing Address in Canada (if different):', '[Student.CanadianAddress]'),
    ('Immigration Status:', '[Student.ResidenceStatus]'),
    ('Citizenship:', '[Student.Citizenship]'),
    ('Gender:', '[Student.Gender]'),
    ('Do you have a study permit?', '[Student.StudyPermit]'),
]

add_simple_table(student_fields)

add_spacer()

# ── Voluntary Disclosure ──
add_heading_styled('VOLUNTARY DISCLOSURE', level=3)

add_static_paragraph('You may voluntarily provide the personal information listed below:')

voluntary_fields = [
    ('Do you identify yourself as an indigenous person (First Nations, Métis, or Inuit)?', '[Student.IndigenousStatus]'),
    ('If you answered "yes", please indicate (First Nations, Métis, or Inuit):', '[Student.IndigenousType]'),
    ('Do you have a long-term physical or mental health condition that limits daily activities?', '[Student fills — Dropdown field in PandaDoc]'),
]

add_simple_table(voluntary_fields)

add_spacer()

# ════════════════════════════════════════════════════════════
# SECTION 2: PROGRAM INFORMATION
# ════════════════════════════════════════════════════════════

add_heading_styled('PROGRAM INFORMATION', level=2)

program_fields = [
    ('Program Title:', '[Program.Title]'),
    ('Hours of Instruction During Contract Term:', '[Program.Hours]'),
    ('Program Duration in Weeks:', '[Program.DurationWeeks]'),
    ('Contract Start Date:', '[Advisor enters — Date field in PandaDoc]'),
    ('Contract End Date:', '[Advisor enters — Date field in PandaDoc]'),
    ('Credential Issued on Graduation:', '[Program.Credential]'),
    ('Program Schedule:', '[Advisor selects — Dropdown: Full-Time / Part-Time]'),
]

add_simple_table(program_fields)

add_spacer()

# Program Delivery Method
add_static_paragraph('Program Delivery Method:', bold=True)

delivery_methods = [
    '☐ In-class',
    '☐ Distance – Synchronous*',
    '☐ Distance – Asynchronous*',
    '☐ Distance – Both Synchronous and Asynchronous*',
    '☐ Combined',
]
for method in delivery_methods:
    add_static_paragraph(f'    {method}')

add_static_paragraph('* Synchronous distance delivery means students attend classes virtually in "real time" with instructors and classmates.', size=8)
add_static_paragraph('* Asynchronous distance learning means students and instructors do not meet in "real time". Students may move through assignments at their own pace.', size=8)

add_spacer()
add_static_paragraph('Language of Instruction: English', bold=True)
add_static_paragraph('Required course materials and technological resources not provided by Western Community College: Students are required to obtain their own equipment for online learning.')

add_spacer()

# ════════════════════════════════════════════════════════════
# SECTION 3: PROGRAM OUTLINE (Smart Content Placeholder)
# ════════════════════════════════════════════════════════════

add_heading_styled('PROGRAM OUTLINE', level=2)

add_static_paragraph('Program Information: Refer to the Attached Program Outline', bold=True)
add_static_paragraph('Admission Requirements: Refer to the Attached Program Outline', bold=True)

add_spacer()

# Smart Content marker
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[SMART CONTENT BLOCK — Program Outline]')
run.font.color.rgb = RGBColor(0, 100, 200)
run.bold = True
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Replace this block with a PandaDoc Smart Content block.\nCondition: IF [Program.Title] equals "[Program Name]" THEN insert [Program Outline Content Library Item]')
run.font.color.rgb = RGBColor(100, 100, 100)
run.italic = True
run.font.size = Pt(9)

add_spacer()

# ════════════════════════════════════════════════════════════
# SECTION 4: WORK EXPERIENCE
# ════════════════════════════════════════════════════════════

add_heading_styled('WORK EXPERIENCE (If applicable)', level=2)
add_static_paragraph('Refer to the Attached Program Outline')

add_spacer()

# ════════════════════════════════════════════════════════════
# SECTION 5: REGULATORY REQUIREMENTS
# ════════════════════════════════════════════════════════════

add_heading_styled('REGULATORY REQUIREMENTS (If applicable)', level=2)
add_static_paragraph('Refer to the Attached Program Outline')

add_spacer()

# ════════════════════════════════════════════════════════════
# SECTION 6: PROGRAM COSTS (Pricing Table)
# ════════════════════════════════════════════════════════════

add_heading_styled('PROGRAM COSTS', level=2)

# Pricing table marker
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[PRICING TABLE BLOCK — Program Fees]')
run.font.color.rgb = RGBColor(0, 100, 200)
run.bold = True
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Replace this with a PandaDoc Pricing Table block with Data Merge enabled.\nLine items will auto-sync from HubSpot deal line items.')
run.font.color.rgb = RGBColor(100, 100, 100)
run.italic = True
run.font.size = Pt(9)

add_spacer()

# Static reference table showing the expected structure
cost_table = doc.add_table(rows=1, cols=4)
cost_table.style = 'Table Grid'
cost_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
headers = ['Item', 'Description', 'Price', 'Subtotal']
for i, header in enumerate(headers):
    cell = cost_table.rows[0].cells[i]
    p = cell.paragraphs[0]
    run = p.add_run(header)
    run.bold = True
    run.font.size = Pt(10)

# Sample rows showing what will auto-populate
sample_rows = [
    ('Application Fee', 'Non-refundable', '$250.00', '$250.00'),
    ('Tuition Fee', 'Based on program + residence status', '[Auto from HubSpot]', '[Auto]'),
    ('Textbook Fee', 'Based on program', '[Auto from HubSpot]', '[Auto]'),
    ('Course Materials Fee', 'If applicable', '[Auto from HubSpot]', '[Auto]'),
    ('Scholarship (if applicable)', 'Advisor enters', '[Advisor enters]', '[Auto]'),
    ('TOTAL PROGRAM COST', '', '', '[Auto-calculated]'),
]

for row_data in sample_rows:
    row = cost_table.add_row()
    for i, val in enumerate(row_data):
        p = row.cells[i].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        if row_data[0] == 'TOTAL PROGRAM COST':
            run.bold = True

add_spacer()

add_static_paragraph('Text Books Fee: Refer to the Attached Book Fee List')
add_static_paragraph('Course Materials Fee: Refer to the Attached Course Material Fee List')

add_spacer()

# ════════════════════════════════════════════════════════════
# SECTION 7: PAYMENT TERMS
# ════════════════════════════════════════════════════════════

add_heading_styled('PAYMENT TERMS', level=2)

add_static_paragraph('Method of Payment:    ☐ Cash    ☐ Cheque    ☐ Credit Card    ☐ Student Loans    ☐ Other')
add_static_paragraph('[Replace checkboxes above with PandaDoc Checkbox fields assigned to Admissions Rep role]', size=8)
add_spacer()
add_static_paragraph('Payment Schedule and/or Financial Plan?    ☐ No    ☐ Yes')
add_static_paragraph('Do you anticipate applying for government student loan?    ☐ No    ☐ Yes')
add_spacer()
add_static_paragraph('* The Admissions process requires students to meet with the College\'s student financial Planner and complete a financial plan prior to admission.', size=8)

add_spacer()

# ════════════════════════════════════════════════════════════
# SECTION 8: STATEMENT OF STUDENT RIGHTS
# ════════════════════════════════════════════════════════════

add_heading_styled('STATEMENT OF STUDENT RIGHTS', level=2)

rights_text = [
    'Western Community College is certified with the Private Training Institutions Regulatory Unit (PTIRU) of the British Columbia Ministry of Post-Secondary Education and Future Skills.',
    '',
    'Before you enrol at a certified private training institution, you should be aware of your rights and responsibilities. You have the right to be treated fairly and respectfully by the institution.',
    '',
    'You have the right to a student enrolment contract that includes the following information:',
    '• amount of tuition and any additional fee for your program refund policy',
    '• if your program includes a work experience, the requirements to participate in the work experience and the geographic area where it will be provided',
    '• whether the program was approved by PTIRU or does not require approval.',
    '',
    'Make sure you read the contract before signing. The institution must provide you with a signed copy.',
    '',
    'You have the right to access the institution\'s dispute resolution process and to be protected against retaliation for making a complaint.',
    '',
    'You have the right to make a claim to PTIRU for a tuition refund if:',
    '• your institution ceased to hold a certificate before you completed an approved program',
    '• you were misled about a significant aspect of your approved program.',
    '',
    'You must file the claim within one year of completing, being dismissed or withdrawing from your program. For more information about PTIRU and how to be an informed student, go to: http://www.privatetraininginstitutions.gov.bc.ca/students/be-an-informed-student',
]

for line in rights_text:
    if line == '':
        add_spacer()
    else:
        add_static_paragraph(line)

add_spacer()

# ════════════════════════════════════════════════════════════
# SECTION 9: REFUND POLICY
# ════════════════════════════════════════════════════════════

add_heading_styled('REFUND POLICY', level=2)

add_static_paragraph('Where applicable, fees will be refunded in accordance with Student Aid BC (SABC) or other provincial loan regulations. Please refer to the College Student Enrolment Contract.')

add_spacer()
add_static_paragraph('Considerations', bold=True)

refund_considerations = 'Western Community College (the College) will refund fees charged for tuition and course materials paid for but not received if the student provides a notice of withdrawal to the College or the College provides a notice of dismissal to the student. Refunds are calculated on the tuition fee of the program. If total fees have not yet been collected, the college is not responsible for refunding more than what has been collected to date and the relevant student may be required to make up for any money due under the contract.'
add_static_paragraph(refund_considerations)

add_static_paragraph('Students are informed of the refund policy during the admission process. Application, administration and assessment fees are non-refundable.')

add_spacer()

# ── Refund Table: In-class/Combined/Synchronous ──
add_static_paragraph('Approved Programs - In-class, Combined Delivery, or Synchronous Distance Delivery', bold=True)

refund_table = doc.add_table(rows=1, cols=2)
refund_table.style = 'Table Grid'

headers_row = refund_table.rows[0]
for i, h in enumerate(['Condition', 'Refund Due']):
    p = headers_row.cells[i].paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(9)

refund_rows = [
    ('No later than 7 days after student signed the enrolment contract, and before the program start date.',
     '100% tuition and all related fees, other than application fee.'),
    ('More than 7 days after the student enrolment contract, and before the program start date.',
     'Institution may retain up to 10% of tuition, to a maximum of $1,000.'),
    ('No later than 7 days after the program start date.',
     'Institution may retain up to 10% of tuition, to a maximum of $1,000.'),
    ('After start date, up to and including 10% of instruction hours.',
     'Institution may retain up to 10% of tuition.'),
    ('After start date, more than 10% but before 30% of instruction hours.',
     'Institution may retain up to 30% of tuition.'),
    ('After start date, more than 30% but before 50% of instruction hours.',
     'Institution may retain up to 50% of tuition.'),
    ('After start date, more than 50% of instruction hours.',
     'No refund due.'),
    ('Student does not attend the first 30% of the program (no-show).',
     'Institution may retain up to 50% of tuition.'),
]

for condition, refund in refund_rows:
    row = refund_table.add_row()
    p0 = row.cells[0].paragraphs[0]
    run0 = p0.add_run(condition)
    run0.font.size = Pt(8)
    p1 = row.cells[1].paragraphs[0]
    run1 = p1.add_run(refund)
    run1.font.size = Pt(8)

add_spacer()

# ── Study Permit Refusal (International) ──
add_static_paragraph('Institution receives a refusal of study permit (international students):', bold=True)
add_static_paragraph('Before 30% of instruction hours would have been provided: 100% tuition and all related fees, other than application fee.')

add_spacer()

# ── Refund Process ──
add_static_paragraph('Process for Refunds', bold=True)
add_static_paragraph('Application fees, student fees, textbooks, fines, dues owing, financial penalties and taxes are non-refundable. Outstanding tuition fees will be deducted from the amount refunded.')
add_static_paragraph('All refund requests must be made in writing. Students must submit the completed and duly signed Refund request form and the supporting documents to the Refund department.')
add_static_paragraph('Refunds required under this policy will be paid to the student, or a person who paid the tuition or fees on behalf of the student, within 30 days.')
add_static_paragraph('An administrative fee of 2% of the total amount paid will be charged for refunds requested by students who paid by credit card.')

add_spacer()

# ── International Student Addendum ──
# Smart Content marker for international-only section
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[SMART CONTENT BLOCK — International Student Refund Addendum]')
run.font.color.rgb = RGBColor(0, 100, 200)
run.bold = True
run.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Show only IF [Student.ResidenceStatus] equals "International Study Permit" OR "Work Permit" OR "Visitor"')
run.font.color.rgb = RGBColor(100, 100, 100)
run.italic = True
run.font.size = Pt(8)

add_spacer()
add_static_paragraph('Additional note for International Students', bold=True)
add_static_paragraph('An international student is a person who is not a Canadian citizen or a landed immigrant or who has been determined under the Immigration Act to be a Convention Refugee. International students require a Study Permit to study in Canada unless they are taking a course or program with a duration of six months or less.')
add_static_paragraph('The College will retain the registration fee due under the enrolment contract for international students who are denied Study Permit authorization from Citizenship and Immigration Canada.')
add_static_paragraph('Students denied a Study Permit must provide the institution with a written request for a refund along with a copy of the denial letter, prior to the program start date given on the institution\'s Letter of Acceptance.')

add_spacer()

# ════════════════════════════════════════════════════════════
# SECTION 10: CONTRACT TERMS & CONDITIONS
# ════════════════════════════════════════════════════════════

add_heading_styled('CONTRACT TERMS AND CONDITIONS', level=2)

# This is a long legal section - included as key clauses
terms_sections = [
    ('Service Nature of Contract / Parties\' Relationship',
     'This is a service contract, and the Services provided by the Institution to the Student – including work experience program components - are as described in this contract. This Contract does not create between the Parties any form of agency, partnership, joint venture or collaboration relationship.'),

    ('Institution Appoints No Third Parties to Act for Them',
     'The Institution has not appointed or authorized any third party to: (a) Enter service and other contracts or incur legal obligations on their behalf. (b) Receive or pay money for them regarding this Contract.'),

    ('The Parties\' Contract Entry and Continuance Conditions',
     'The Institution represents that its admission standards existing pre-contract and described and restated in this Contract were applied fairly, impartially and objectively and without waiver by the Institution. It further represents that its decision to admit the student was made because the Student was reasonably determined by the Institution to meet Institution admission standards and capable of succeeding in and graduating from the Program in the set contract period.'),

    ('Institution Contract Entry Conditions',
     'The Institution enters and continues as a Contract Party on the following conditions:\n1. The Student represented and represents that their application information submitted to establish they meet Institution and Program admissions standards was and is true and complete.\n2. The Student met and meets admission standards existing pre-contract and restated in this Contract.\n3. The Student will immediately inform the Institution of changes to application information.'),

    ('Student Contract Entry Conditions',
     'The Student enters and continues as a Contract Party on the following conditions:\n• The Service Program the Institution will provide to the Student under this Contract will have the elements and features represented pre-contract and promised in this Contract by the Institution.\n• The Institution and Service Program were and are compliant with the laws and regulations described in this Contract.'),

    ('Services to Help Student Achieve Learning Objective',
     'The Institution will do all that it reasonably can to help the Students meet progress and graduation standards set in this Contract in the contract set period, but it does not promise or guarantee the Student they will achieve the service program learning objective or graduate in the contract set period.'),

    ('Assessment Systems and Services',
     'The Institution will provide fair, impartial and objective assessment systems and services that determine if the Student meets academic standards and can progress in or graduate from the Service Program. The Institution will keep the Student informed of their progress and the likelihood they will graduate.'),

    ('Transcript Systems and Services',
     'The Institution will assemble and publish in print and digital form a final transcript that records the Student\'s work and achievement at the Institution and in the Program. No fees will be charged to students for the initial final transcript.'),

    ('No Promise of Credit Transfer',
     'The Institution makes no representation or promise that the Services it provides to the Student under this Contract are uniform with or comparable to those provided by other institutions or parties. The Institution makes no representation or promise that any credit or credential granted to the Student under this Contract will be accepted or recognized by other institutions or parties.'),

    ('Unpromised Post Contract Services or Outcomes',
     'The Institution makes no promise in this Contract that the Student will, during or after this Contract, receive from a third party specific: business, workplace or other contract offers, opportunities or results; financial incomes or compensation; credit or credential recognitions, immigration status or professional memberships.'),

    ('Contract Service Period and Termination',
     'The effective date of this Contract will be when all Parties assent to it, or a date set in this Contract. The Institution may by written notice terminate this Contract if it reasonably determines that the Student: knowingly provided incomplete or untrue application information; is guilty of academic misconduct or dishonesty; is guilty of non-academic misconduct that has or reasonably may have adverse impacts on Institution functions or on the health, safety, rights, or property of the Institution or its staff and students.'),

    ('Work Experience Programs',
     'If the Institution establishes a work experience program element, the Institution will solely provide and be legally responsible for provision of the primary Service Program and its Work Experience Program component to the Student. The Workplace Host will not be party to this Contract.'),

    ('Intellectual Property',
     'The Institution holds rights to its logos, product names, copyrighted works, patents, industrial designs, trademarks, curricula and program materials and grants to the Student a limited, non-transferable and non-exclusive license to use them during the contract period.'),

    ('Personal Information Management',
     'The collection, management and dissemination of personal information pursuant to this Contract will comply with all applicable laws and the provisions of this Contract. The Institution will only use the Student\'s personal information collected by the Institution for the purpose for which it was collected or for uses consistent with that purpose.'),
]

for title, content in terms_sections:
    add_static_paragraph(title, bold=True, size=10)
    add_static_paragraph(content, size=9)
    add_spacer()


# ════════════════════════════════════════════════════════════
# SECTION 11: ANCESTRAL TERRITORY ACKNOWLEDGMENT
# ════════════════════════════════════════════════════════════

add_heading_styled('ANCESTRAL TERRITORY RECOGNITION AND ACKNOWLEDGMENT', level=2)

add_static_paragraph('Western Community College recognizes and acknowledges the ancestral unceded territory of the Coastal First Nations (Sumas First Nation, Mastqui First Nation, Semiahmoo, Katzie, Kwikwetlem, Kwantlen, Qayqayt and Tsawwassen) where we live and work. We honour these First Nations for sharing and accepting our families and students in their Traditional territories.')

add_spacer()

# ════════════════════════════════════════════════════════════
# SECTION 12: PTIRU & PRIVACY STATEMENTS
# ════════════════════════════════════════════════════════════

add_heading_styled('PRIVATE TRAINING INSTITUTIONS REGULATORY UNIT (PTIRU)', level=2)

add_static_paragraph('This institution is certified by the Private Training Institutions Regulatory Unit (PTIRU) of the Ministry of Post-Secondary Education and Future Skills. Certified institutions must comply with regulatory requirements relating to, among other things, student enrolment contracts, tuition refunds and instructor qualifications. For more information about PTIRU, go to www.privatetraininginstitutions.gov.bc.ca.')

add_spacer()

add_static_paragraph('Please be advised that under section 61 of the Private Training Act, the Registrar is authorized to collect, use and disclose personal information in accordance with the Registrar\'s regulatory duties under that Act. Accordingly, this institution is authorized to disclose your personal information to the Registrar for regulatory purposes.')

add_spacer()

# ── International Student Consent (Conditional) ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[SMART CONTENT BLOCK — International Student Consent]')
run.font.color.rgb = RGBColor(0, 100, 200)
run.bold = True
run.font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Show only IF [Student.ResidenceStatus] is International')
run.font.color.rgb = RGBColor(100, 100, 100)
run.italic = True
run.font.size = Pt(8)

add_spacer()

add_static_paragraph('I consent to the sharing, in accordance with applicable Provincial privacy legislation, of my enrollment and reporting information between this institution and Immigration, Refugees and Citizenship Canada, as necessary, for the purposes of the International Student Program.')

add_spacer()

# ════════════════════════════════════════════════════════════
# SECTION 13: STUDENT DECLARATION
# ════════════════════════════════════════════════════════════

add_heading_styled('STUDENT DECLARATION', level=2)

add_static_paragraph('I consent to the institution sharing my personal information with the Ministry of Post-Secondary Education and Future Skills for research purposes and statistical analysis under the authority of sections 6(2)(a) and 10(1)(a) of the Personal Information Protection Act (PIPA).')

add_spacer()

add_static_paragraph('I consent to the institution sharing my personal information with Immigration, Refugees and Citizenship Canada for the purposes of the International Student Program under the authority of section 6(2)(a) and 10(1)(a) of the Personal Information Protection Act (PIPA).')

add_spacer()

add_static_paragraph('Should you have any questions about the collection, disclosure and use of personal information you may contact: Director, Policy and Institution Certification, Private Training Institutions Regulatory Unit, Governance, Legislation and Corporate Planning Division, Ministry of Post-Secondary Education and Future Skills, 310-601 Cordova Street W, Vancouver, BC V6B 1G1 or by telephone at (604) 569-0019.', size=8)

add_spacer()

add_static_paragraph('By signing below, I certify the following:', bold=True)

declarations = [
    'All information/statements on this application form and supporting documents are true and complete. I authorize Western Community College to verify any information provided as part of this application. I understand that the registration fee/application fee is non-refundable.',
    'I understand that evidence of falsified documents is shared with other Canadian colleges and universities.',
    'I understand and acknowledge that it is my responsibility to be aware of, and comply with all WCC policies and procedures. Admission is subject to assessment of qualifications and availability of seats.',
]

for decl in declarations:
    p = doc.add_paragraph()
    run = p.add_run(f'• {decl}')
    run.font.size = Pt(9)

add_spacer()

# ════════════════════════════════════════════════════════════
# SECTION 14: SIGNATURES
# ════════════════════════════════════════════════════════════

add_heading_styled('SIGNATURES', level=2)

# Student Signature Block
add_static_paragraph('STUDENT SIGNATURE', bold=True, size=12)

sig_table_student = doc.add_table(rows=2, cols=2)
sig_table_student.style = 'Table Grid'

# Row 1
sig_table_student.rows[0].cells[0].paragraphs[0].add_run('Student Signature:').bold = True
p = sig_table_student.rows[0].cells[1].paragraphs[0]
run = p.add_run('[Signature field — Student role]')
run.font.color.rgb = RGBColor(0, 100, 200)

# Row 2
sig_table_student.rows[1].cells[0].paragraphs[0].add_run('Date Signed:').bold = True
p = sig_table_student.rows[1].cells[1].paragraphs[0]
run = p.add_run('[Date field — Student role]')
run.font.color.rgb = RGBColor(0, 100, 200)

add_spacer()
add_spacer()

# Institution Signature Block
add_static_paragraph('INSTITUTION SIGNATURE', bold=True, size=12)

sig_table_inst = doc.add_table(rows=2, cols=2)
sig_table_inst.style = 'Table Grid'

# Row 1
sig_table_inst.rows[0].cells[0].paragraphs[0].add_run('Signature of Institution Representative:').bold = True
p = sig_table_inst.rows[0].cells[1].paragraphs[0]
run = p.add_run('[Signature field — Admissions Rep role]')
run.font.color.rgb = RGBColor(0, 100, 200)

# Row 2
sig_table_inst.rows[1].cells[0].paragraphs[0].add_run('Date Signed:').bold = True
p = sig_table_inst.rows[1].cells[1].paragraphs[0]
run = p.add_run('[Date field — Admissions Rep role]')
run.font.color.rgb = RGBColor(0, 100, 200)


# ════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════

output_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(output_dir, '..', 'templates', 'WCC_Master_Student_Enrollment_Contract.docx')
doc.save(output_path)
print(f'Template saved to: {os.path.abspath(output_path)}')
print()
print('=== TEMPLATE BUILDING INSTRUCTIONS ===')
print()
print('After uploading this DOCX to PandaDoc as a template:')
print()
print('1. TOKENS: All [Token.Name] placeholders are ready for PandaDoc variables.')
print('   In PandaDoc editor, replace bracketed text with actual PandaDoc variables.')
print()
print('2. ROLES: Create two roles:')
print('   - Student (signing order 1)')
print('   - Admissions Rep (signing order 2)')
print()
print('3. FIELDS TO ADD IN PANDADOC EDITOR:')
print('   - Signature field (Student role)')
print('   - Signature field (Admissions Rep role)')
print('   - Date field x2 (one per role)')
print('   - PEN text field (Admissions Rep role)')
print('   - Contract Start Date (Admissions Rep role)')
print('   - Contract End Date (Admissions Rep role)')
print('   - Program Schedule dropdown: Full-Time / Part-Time (Admissions Rep role)')
print('   - Delivery Method checkboxes (Admissions Rep role)')
print('   - Payment Method checkboxes (Admissions Rep role)')
print('   - Disability disclosure dropdown (Student role)')
print()
print('4. PRICING TABLE:')
print('   - Replace the sample fee table with a PandaDoc Pricing Table block')
print('   - Enable Data Merge')
print('   - Columns: Name, Description, Price, QTY, Subtotal')
print('   - Add a discount row for Scholarship')
print()
print('5. SMART CONTENT BLOCKS (3 total):')
print('   Block 1: Program Outline — 65 conditions on [Program.Title]')
print('   Block 2: International Refund Addendum — condition on [Student.ResidenceStatus]')
print('   Block 3: International Student Consent — condition on [Student.ResidenceStatus]')
print()
print('6. CONTENT LIBRARY ITEMS TO CREATE:')
print('   - 65 Program Outline items (one per program)')
print('   - Optional: Static sections as library items for easy maintenance')
