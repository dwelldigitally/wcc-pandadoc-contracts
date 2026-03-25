"""
Generate DOCX files for the 4 static Content Library items,
then upload them to PandaDoc via API.
"""
import subprocess, json, os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'libraries')

def styled_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    return doc

def add_text(doc, text, bold=False, size=10, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Arial'
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def save_and_report(doc, filename):
    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    print(f"  Created: {path}")
    return path


# ═══════════════════════════════════════════
# 1. STATEMENT OF STUDENT RIGHTS
# ═══════════════════════════════════════════
print("1. Statement of Student Rights")
doc = styled_doc()

add_text(doc, 'STATEMENT OF STUDENT RIGHTS', bold=True, size=12, color=(139, 0, 0))
doc.add_paragraph()

texts = [
    'Western Community College is certified with the Private Training Institutions Regulatory Unit (PTIRU) of the British Columbia Ministry of Post-Secondary Education and Future Skills.',
    'Before you enrol at a certified private training institution, you should be aware of your rights and responsibilities. You have the right to be treated fairly and respectfully by the institution.',
    'You have the right to a student enrolment contract that includes the following information:',
]
for t in texts:
    add_text(doc, t)

bullets = [
    'amount of tuition and any additional fee for your program refund policy',
    'if your program includes a work experience, the requirements to participate in the work experience and the geographic area where it will be provided',
    'whether the program was approved by PTIRU or does not require approval.',
]
for b in bullets:
    p = doc.add_paragraph(f'• {b}')
    p.style.font.size = Pt(10)

doc.add_paragraph()
add_text(doc, 'Make sure you read the contract before signing. The institution must provide you with a signed copy.')
doc.add_paragraph()
add_text(doc, 'You have the right to access the institution\'s dispute resolution process and to be protected against retaliation for making a complaint.')
doc.add_paragraph()
add_text(doc, 'You have the right to make a claim to PTIRU for a tuition refund if:')

claims = [
    'your institution ceased to hold a certificate before you completed an approved program',
    'you were misled about a significant aspect of your approved program.',
]
for c in claims:
    doc.add_paragraph(f'• {c}')

doc.add_paragraph()
add_text(doc, 'You must file the claim within one year of completing, being dismissed or withdrawing from your program. For more information about PTIRU and how to be an informed student, go to: http://www.privatetraininginstitutions.gov.bc.ca/students/be-an-informed-student')

path1 = save_and_report(doc, 'Statement_of_Student_Rights.docx')


# ═══════════════════════════════════════════
# 2. CONTRACT TERMS AND CONDITIONS
# ═══════════════════════════════════════════
print("2. Contract Terms and Conditions")
doc = styled_doc()

add_text(doc, 'CONTRACT TERMS AND CONDITIONS', bold=True, size=12, color=(139, 0, 0))
doc.add_paragraph()

terms = [
    ('Service Nature of Contract / Parties\' Relationship',
     'This is a service contract, and the Services provided by the Institution to the Student \u2013 including work experience program components - are as described in this contract. This Contract does not create between the Parties any form of agency, partnership, joint venture or collaboration relationship.'),
    ('Institution Appoints No Third Parties to Act for Them',
     'The Institution has not appointed or authorized any third party to: (a) Enter service and other contracts or incur legal obligations on their behalf. (b) Receive or pay money for them regarding this Contract. If the Student pays money to a third party under a Student and third-party contract that requires the third party to transmit money to the Institution for Student payments under this Contract, the Institution will not be party to or responsible for third-party actions under that Student and third-party contract.'),
    ('Admission Decision',
     'The Institution represents that its admission standards existing pre-contract and described and restated in this Contract were applied fairly, impartially and objectively and without waiver by the Institution. It further represents that its decision to admit the student was made because the Student was reasonably determined by the Institution \u2013 given the Student\u2019s application information - to meet Institution admission standards and capable of succeeding in and graduating from the Program in the set contract period.'),
    ('Institution Contract Entry Conditions',
     'The Institution enters and continues as a Contract Party on the following conditions:\n1. The Student represented and represents that their application information submitted to establish they meet Institution and Program admissions standards was and is true and complete.\n2. The Student met and meets admission standards existing pre-contract and restated in this Contract.\n3. The Student will immediately inform the Institution of changes to application information.'),
    ('Pre-Requisites for Later Elements of Program',
     'If service program elements are pre-requisites to taking later elements in the Service Program, the rights and duties of the Parties related to later elements are conditional and subject to completion of pre-requisites. The Student will be admitted to later program elements if pre-requisite conditions are met.'),
    ('Student Contract Entry Conditions',
     'The Student enters and continues as a Contract Party on the following conditions:\n\u2022 The Service Program the Institution will provide to the Student under this Contract will have the elements and features represented pre-contract and promised in this Contract by the Institution.\n\u2022 The Institution and Service Program were and are compliant with the laws and regulations described in this Contract.\n\u2022 The Institution provides the Program promised in this Contract \u2013 subject to Contract termination.'),
    ('International Students \u2013 Study Permits',
     'If the Student requires an immigration permit to enter and study in the jurisdiction where the Institution provides onsite services, this Contract is conditional on the Student receiving and having the permit.'),
    ('Contract Inclusions \u2013 Entire Agreement',
     'Subject to its express inclusion provisions, this Contract constitutes the entire agreement between the Parties and supersedes and does not include all other verbal or written communications and agreements not otherwise intentionally and expressly incorporated by its Parties into this Contract.'),
    ('Services to Help Student Achieve Learning Objective',
     'The Institution will do all that it reasonably can to help the Students meet progress and graduation standards set in this Contract in the contract set period, but it does not promise or guarantee the Student they will achieve the service program learning objective or graduate in the contract set period. The Student must inform the Institution if they need additional services to help them meet learning objectives in contract set periods.'),
    ('Service Program Graduates \u2013 Regulated Profession',
     'If the service program can lead to employment in a career occupation governed by a profession or vocation group regulator and inclusion in the group is required before work can be performed, that requirement will be set out in service program outlines. Service program graduation does not guarantee membership in the profession or vocational group.'),
    ('Assessment Systems and Services',
     'The Institution will provide fair, impartial and objective assessment systems and services that determine if the Student meets academic standards and can progress in or graduate from the Service Program. The Institution will keep the Student informed of their progress and the likelihood they will graduate.'),
    ('Transcript Systems and Services',
     'The Institution will assemble and publish in print and digital form a final transcript that records the Student\u2019s work and achievement at the Institution and in the Program \u2013 and provide it to the Student at their request. No fees will be charged to students for the initial final transcript.'),
    ('No Promise of Credit Transfer',
     'The Institution makes no representation or promise that the Services it provides to the Student under this Contract are uniform with or comparable to those provided by other institutions or parties. The Institution makes no representation or promise that any credit or credential granted to the Student under this Contract will be accepted or recognized by other institutions or parties.'),
    ('Unpromised Post Contract Services or Outcomes',
     'The Institution makes no promise in this Contract that the Student will, during or after this Contract, receive from a third party specific: business, workplace or other contract offers, opportunities or results; financial incomes or compensation; credit or credential recognitions, immigration status or professional memberships.'),
    ('Contract Service Period and Termination',
     'The effective date of this Contract will be when all Parties assent to it, or a date set in this Contract. The Institution may by written notice terminate this Contract if it reasonably determines that the Student: knowingly provided incomplete or untrue application information; is guilty of academic misconduct or dishonesty; is guilty of non-academic misconduct that has or reasonably may have adverse impacts on Institution functions or on the health, safety, rights, or property of the Institution or its staff and students.'),
    ('Work Experience Programs',
     'If the Institution establishes a work experience program element involving the Student studying or working with, for or under supervision of a third party that is not party to this Contract, that work program element will be part of the Service provided to the Student in this Contract only if this Contract expressly provides for that inclusion. The Institution will solely provide and be legally responsible for provision of the primary Service Program and its Work Experience Program component to the Student. The Workplace Host will not be party to this Contract.'),
    ('Student Responsibilities During Work Experience',
     'The Student will: Be punctual and prepared for work experience days and periods. Be familiar with host sites, policies, procedures, routines and staff. Perform specific tasks assigned in a skillful, orderly, organized, safe and professional manner. Meet learning goals for the work experience program. Protect the intellectual property and confidential information of the Institution and host. Accept constructive feedback regarding their work performance.'),
    ('Intellectual Property',
     'The Institution holds rights to its logos, product names, copyrighted works, patents, industrial designs, trademarks, curricula and program materials and grants to the Student a limited, non-transferable and non-exclusive license to use them during the contract period. The Student will at all times protect and not harm the Institution\u2019s interests and rights.'),
    ('Personal Information Management',
     'The collection, management and dissemination of personal information pursuant to this Contract will comply with all applicable laws and the provisions of this Contract. The Institution will only use the Student\u2019s personal information collected by the Institution for the purpose for which it was collected or for uses consistent with that purpose, unless the Student consents otherwise.'),
]

for title, content in terms:
    add_text(doc, title, bold=True, size=10)
    add_text(doc, content, size=9)
    doc.add_paragraph()

path2 = save_and_report(doc, 'Contract_Terms_and_Conditions.docx')


# ═══════════════════════════════════════════
# 3. STUDENT DECLARATION
# ═══════════════════════════════════════════
print("3. Student Declaration")
doc = styled_doc()

add_text(doc, 'STUDENT DECLARATION', bold=True, size=12, color=(139, 0, 0))
doc.add_paragraph()

add_text(doc, 'I consent to the institution sharing my personal information with the Ministry of Post-Secondary Education and Future Skills for research purposes and statistical analysis under the authority of sections 6(2)(a) and 10(1)(a) of the Personal Information Protection Act (PIPA).')
doc.add_paragraph()
add_text(doc, 'I consent to the institution sharing my personal information with Immigration, Refugees and Citizenship Canada for the purposes of the International Student Program under the authority of section 6(2)(a) and 10(1)(a) of the Personal Information Protection Act (PIPA).')
doc.add_paragraph()
add_text(doc, 'Should you have any questions about the collection, disclosure and use of personal information you may contact: Director, Policy and Institution Certification, Private Training Institutions Regulatory Unit, Governance, Legislation and Corporate Planning Division, Ministry of Post-Secondary Education and Future Skills, 310-601 Cordova Street W, Vancouver, BC V6B 1G1 or by telephone at (604) 569-0019.', size=8)
doc.add_paragraph()
add_text(doc, 'By signing below, I certify the following:', bold=True)

declarations = [
    'All information/statements on this application form and supporting documents are true and complete. I authorize Western Community College to verify any information provided as part of this application. I understand that the registration fee/application fee is non-refundable.',
    'I understand that evidence of falsified documents is shared with other Canadian colleges and universities.',
    'I understand and acknowledge that it is my responsibility to be aware of, and comply with all WCC policies and procedures. Admission is subject to assessment of qualifications and availability of seats.',
    'I consent to the sharing, in accordance with applicable Provincial privacy legislation, of my enrolment and reporting information between Western Community College and Immigration, Refugees and Citizenship Canada, as necessary, for the purposes of the International Student Program.',
]
for d in declarations:
    doc.add_paragraph(f'\u2022 {d}')

path3 = save_and_report(doc, 'Student_Declaration.docx')


# ═══════════════════════════════════════════
# 4. PTIRU STATEMENT
# ═══════════════════════════════════════════
print("4. PTIRU Statement")
doc = styled_doc()

add_text(doc, 'PRIVATE TRAINING INSTITUTIONS REGULATORY UNIT (PTIRU)', bold=True, size=12, color=(139, 0, 0))
doc.add_paragraph()

add_text(doc, 'This institution is certified by the Private Training Institutions Regulatory Unit (PTIRU) of the Ministry of Post-Secondary Education and Future Skills. Certified institutions must comply with regulatory requirements relating to, among other things, student enrolment contracts, tuition refunds and instructor qualifications. For more information about PTIRU, go to www.privatetraininginstitutions.gov.bc.ca.')
doc.add_paragraph()
add_text(doc, 'Please be advised that under section 61 of the Private Training Act, the Registrar is authorized to collect, use and disclose personal information in accordance with the Registrar\u2019s regulatory duties under that Act. Accordingly, this institution is authorized to disclose your personal information to the Registrar for regulatory purposes.')
doc.add_paragraph()
add_text(doc, 'I consent to the sharing, in accordance with applicable Provincial privacy legislation, of my enrollment and reporting information between this institution and Immigration, Refugees and Citizenship Canada, as necessary, for the purposes of the International Student Program.')

path4 = save_and_report(doc, 'PTIRU_Statement.docx')

print()
print("All 4 DOCX files created. Now uploading to PandaDoc...")
print()

# ═══════════════════════════════════════════
# UPLOAD TO PANDADOC
# ═══════════════════════════════════════════

import os
API_KEY = os.environ.get("PANDADOC_API_KEY", "")
items = [
    ("snfvad5LqzsY2uJnhqgUtE", path1, "Statement of Student Rights"),
    ("3Sx4WPhrSMgwmkhup2NZLZ", path2, "Contract Terms and Conditions"),
    ("dXqxi2R3JZZMoJzVFtoRSj", path3, "Student Declaration"),
    ("hXQubJ6usUmNnFwf3GCPbN", path4, "PTIRU Statement"),
]

for item_id, filepath, name in items:
    print(f"  Uploading: {name}...")
    result = subprocess.run(
        [
            "curl", "-s", "-X", "PATCH",
            f"https://api.pandadoc.com/public/v1/content-library-items/{item_id}",
            "-H", f"Authorization: API-Key {API_KEY}",
            "-F", f"file=@{filepath}",
        ],
        capture_output=True, text=True
    )
    try:
        data = json.loads(result.stdout)
        print(f"    Status: {data.get('status', 'unknown')} | ID: {data.get('id', 'error')}")
    except:
        print(f"    Response: {result.stdout[:200]}")

print()
print("Done! All 4 content library items populated.")
