"""
WCC Contract Generator v2 — Pre-filled DOCX approach

Fills the Word template with student/program/fee data, then uploads
the completed DOCX to PandaDoc for signature collection only.

Usage:
    python generate-contract-v2.py <hubspot-deal-id> [intake-date]

    intake-date: Optional YYYY-MM-DD. If provided, uses this specific intake
                 and matches fees by effective_from date. If omitted, picks
                 the next open intake automatically.

Environment variables:
    PANDADOC_API_KEY    PandaDoc API key
    HUBSPOT_TOKEN       HubSpot private app token
    GOOGLE_SHEETS_ID    Google Sheet ID (optional — falls back to local CSV)
    GOOGLE_API_KEY      Google API key (optional)
"""

import sys
import os
import csv
import json
import copy
import re
import time
import tempfile
import urllib.request
import urllib.error
from pathlib import Path
from datetime import datetime, timezone, timedelta

from lxml import etree

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from fee_matching import match_fees, resolve_fee_amounts, is_domestic, DOMESTIC_STATUSES

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

# Pacific Daylight Time (UTC-7) — BC is in PDT most of the year
PST = timezone(timedelta(hours=-7))

SCRIPT_DIR = Path(__file__).parent
DATA_DIR = SCRIPT_DIR.parent / "data"
TEMPLATE_PATH = SCRIPT_DIR.parent / "New Contract Format - Final.docx"
OUTPUT_DIR = SCRIPT_DIR.parent / "output"

def require_env(var_name):
    """Return the value of an environment variable or exit with a clear error."""
    value = os.environ.get(var_name)
    if not value:
        print(f"ERROR: Required environment variable {var_name} is not set.")
        sys.exit(1)
    return value


PANDADOC_API_KEY = require_env("PANDADOC_API_KEY")
PANDADOC_BASE_URL = "https://api.pandadoc.com/public/v1"

HUBSPOT_TOKEN = require_env("HUBSPOT_TOKEN")
HUBSPOT_BASE_URL = "https://api.hubapi.com"

GOOGLE_SHEETS_ID = os.environ.get("GOOGLE_SHEETS_ID", "")
GOOGLE_API_KEY = os.environ.get("GOOGLE_API_KEY", "")

PANDADOC_STATUS_DRAFT = "document.draft"
PANDADOC_STATUS_ERROR_PREFIX = "error"

INSTITUTION = {
    "name": "Western Community College Inc.",
    "address": "#201 8313 120th Street, Surrey, BC V3W 3N4",
    "number": "3758",
}


# ---------------------------------------------------------------------------
# HTTP helpers
# ---------------------------------------------------------------------------

def api_get(url, token, token_type="Bearer"):
    req = urllib.request.Request(url)
    req.add_header("Authorization", f"{token_type} {token}")
    req.add_header("Content-Type", "application/json")
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            return json.loads(resp.read().decode())
    except urllib.error.HTTPError as e:
        err_body = e.read().decode() if e.fp else ""
        raise Exception(f"GET {url} failed — HTTP {e.code}: {err_body}") from e


def api_post(url, data, token, token_type="Bearer", content_type="application/json"):
    body = json.dumps(data).encode() if isinstance(data, dict) else data
    req = urllib.request.Request(url, data=body, method="POST")
    req.add_header("Authorization", f"{token_type} {token}")
    req.add_header("Content-Type", content_type)
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            return json.loads(resp.read().decode())
    except urllib.error.HTTPError as e:
        err_body = e.read().decode() if e.fp else ""
        raise Exception(f"HTTP {e.code}: {err_body}") from e


def api_post_multipart(url, fields, files, token, token_type="API-Key"):
    """Multipart form upload for PandaDoc document creation."""
    import uuid
    boundary = f"----WCCBoundary{uuid.uuid4().hex}"
    body = b""

    for key, value in fields.items():
        body += f"--{boundary}\r\n".encode()
        body += f'Content-Disposition: form-data; name="{key}"\r\n\r\n'.encode()
        body += f"{value}\r\n".encode()

    for key, (filename, filedata, mime) in files.items():
        body += f"--{boundary}\r\n".encode()
        body += f'Content-Disposition: form-data; name="{key}"; filename="{filename}"\r\n'.encode()
        body += f"Content-Type: {mime}\r\n\r\n".encode()
        body += filedata
        body += b"\r\n"

    body += f"--{boundary}--\r\n".encode()

    req = urllib.request.Request(url, data=body, method="POST")
    req.add_header("Authorization", f"{token_type} {token}")
    req.add_header("Content-Type", f"multipart/form-data; boundary={boundary}")
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            return json.loads(resp.read().decode())
    except urllib.error.HTTPError as e:
        err_body = e.read().decode() if e.fp else ""
        raise Exception(f"HTTP {e.code}: {err_body}") from e


# ---------------------------------------------------------------------------
# Data loading — Google Sheets or local CSV
# ---------------------------------------------------------------------------

def parse_csv(filepath):
    rows = []
    with open(filepath, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append({k.strip(): v.strip() for k, v in row.items()})
    return rows


def load_local_data():
    programs = parse_csv(DATA_DIR / "demo-programs.csv")
    fees = parse_csv(DATA_DIR / "demo-fees.csv")
    return {"programs": programs, "fees": fees, "intakes": []}


def load_sheet_data():
    """Load from Google Sheets if configured."""
    if not GOOGLE_SHEETS_ID or not GOOGLE_API_KEY:
        return None
    try:
        result = {}
        for tab in ["Programs", "Fees", "Intakes"]:
            url = (
                f"https://sheets.googleapis.com/v4/spreadsheets/{GOOGLE_SHEETS_ID}"
                f"/values/{tab}?key={GOOGLE_API_KEY}"
            )
            data = api_get(url, "", token_type="")
            headers = data["values"][0]
            rows = []
            for row in data["values"][1:]:
                obj = {}
                for i, h in enumerate(headers):
                    obj[h.strip()] = row[i].strip() if i < len(row) else ""
                rows.append(obj)
            result[tab.lower()] = rows
        return result
    except Exception as e:
        print(f"  Google Sheets failed ({e}), falling back to local CSV...")
        return None


def load_data():
    data = load_sheet_data()
    if data:
        print("  -> Loaded from Google Sheets")
        return data
    print("  -> Using local CSV files")
    return load_local_data()


# ---------------------------------------------------------------------------
# HubSpot
# ---------------------------------------------------------------------------

def get_deal_data(deal_id):
    deal_props = "dealname,program_of_study,program_hours,program_duration_weeks,program_credential,hubspot_owner_id"
    deal = api_get(
        f"{HUBSPOT_BASE_URL}/crm/v3/objects/deals/{deal_id}?properties={deal_props}",
        HUBSPOT_TOKEN,
    )

    assoc = api_get(
        f"{HUBSPOT_BASE_URL}/crm/v3/objects/deals/{deal_id}/associations/contacts",
        HUBSPOT_TOKEN,
    )
    if not assoc.get("results"):
        raise Exception(f"No contact associated with deal {deal_id}")
    contact_id = assoc["results"][0]["id"]

    contact_props = (
        "firstname,lastname,email,phone,mobilephone,date_of_birth,"
        "address,city,state,zip,country,"
        "program_of_study,residence_status,citizenship,gender"
    )
    contact = api_get(
        f"{HUBSPOT_BASE_URL}/crm/v3/objects/contacts/{contact_id}?properties={contact_props}",
        HUBSPOT_TOKEN,
    )

    # Get deal owner (advisor) details
    owner_id = deal["properties"].get("hubspot_owner_id")
    advisor = {"firstName": "Admissions", "lastName": "Representative", "email": "admissions@wcc.ca"}
    if owner_id:
        try:
            advisor = api_get(
                f"{HUBSPOT_BASE_URL}/crm/v3/owners/{owner_id}",
                HUBSPOT_TOKEN,
            )
        except Exception:
            pass  # Fall back to default

    return {
        "deal": deal["properties"],
        "contact": contact["properties"],
        "advisor": advisor,
        "deal_id": deal_id,
        "contact_id": contact_id,
    }


# ---------------------------------------------------------------------------
# Lookups
# ---------------------------------------------------------------------------

def lookup_program(programs, name):
    for p in programs:
        if p["program_name"].lower() == name.lower():
            return p
    return None


def lookup_next_intake(intakes, name):
    now = datetime.now(PST)
    candidates = [
        i for i in intakes
        if i.get("program_name", "").lower() == name.lower()
        and i.get("status") == "Open"
        and i.get("intake_date")
    ]
    candidates.sort(key=lambda x: x["intake_date"])
    for c in candidates:
        try:
            if datetime.strptime(c["intake_date"], "%Y-%m-%d").replace(tzinfo=PST) >= now:
                return c
        except ValueError:
            pass
    return None


# ---------------------------------------------------------------------------
# DOCX Manipulation — Uses "New Contract Format - Final.docx"
#
# Strategy: The original doc has label cells with an empty Para 1 for values.
# We write the value into that empty paragraph, matching the existing font
# (Times New Roman 10pt). For fees, we deep-copy rows to preserve all XML.
# Everything else (borders, shading, colors, images, sections) is untouched.
# ---------------------------------------------------------------------------

from docx.shared import Pt as DocxPt
from docx.oxml.ns import qn


def replace_sdt_with_plain_text(sdt, value):
    """Remove a structured document tag (dropdown/content control) from the XML
    and replace it with a plain text run, preserving the formatting of the
    original run inside the SDT.

    This completely removes the dropdown — PandaDoc will see regular text only.
    """
    value = value or ""

    # Get the parent element
    parent = sdt.getparent()
    if parent is None:
        return

    # Find the SDT content element and extract formatting from its first run
    sdt_content = sdt.find(qn('w:sdtContent'))
    if sdt_content is None:
        return

    # Find the first run inside the SDT content
    first_run = sdt_content.find('.//' + qn('w:r'))
    if first_run is not None:
        # Clone the run, replace its text, and insert it where the SDT was
        new_run = copy.deepcopy(first_run)
        # Clear all text elements and set the new value
        for t in new_run.findall(qn('w:t')):
            t.getparent().remove(t)
        # Add new text element
        t_elem = etree.SubElement(new_run, qn('w:t'))
        t_elem.text = value
        t_elem.set(qn('xml:space'), 'preserve')

        # Find the paragraph containing this SDT
        sdt_para = sdt.getparent()
        if sdt_para is not None:
            # Insert the run where the SDT was, then remove the SDT
            idx = list(sdt_para).index(sdt)
            sdt_para.insert(idx, new_run)
            sdt_para.remove(sdt)
    else:
        # No run found — just remove the SDT entirely
        parent.remove(sdt)


def replace_all_dropdowns(doc, contact, program, intake):
    """Find ALL dropdown/combobox content controls in the document and
    replace them with plain text values. Checkboxes are left alone.
    """
    body = doc.element.body
    sdts = body.findall('.//' + qn('w:sdt'))

    schedule = intake.get("schedule", "") if intake else ""
    residence = contact.get("residence_status") or ""
    has_study_permit = "Yes" if "study" in residence.lower() or "permit" in residence.lower() else "No"

    # Map SDT index -> value (based on order found in document)
    sdt_values = {
        0: residence,                               # Immigration Status
        1: contact.get("gender") or "",              # Gender
        2: has_study_permit,                         # Study permit?
        3: "",                                       # Indigenous — leave blank
        4: "",                                       # Health condition — leave blank
        5: program.get("credential") or "",           # Credential
        6: schedule,                                  # Schedule
        # 7-15 are checkboxes — skip
    }

    # Process in REVERSE order so indices don't shift as we remove elements
    for i in sorted(sdt_values.keys(), reverse=True):
        if i < len(sdts):
            sdt = sdts[i]
            pr = sdt.find(qn('w:sdtPr'))
            # Only replace dropdowns/comboboxes, not checkboxes
            is_dropdown = pr is not None and (
                pr.find(qn('w:dropDownList')) is not None or
                pr.find(qn('w:comboBox')) is not None
            )
            if is_dropdown:
                replace_sdt_with_plain_text(sdt, sdt_values[i])


def write_value_to_cell(cell, value, para_index=1, font_name="Times New Roman", font_size=127000):
    """Write a value into a specific paragraph of a cell.

    If the target paragraph exists and is empty, adds a run with matching font.
    If the target paragraph doesn't exist, ADDS a new paragraph (never overwrites
    the label in para 0).
    """
    value = value or ""

    if para_index < len(cell.paragraphs):
        # Target paragraph exists — write into it
        para = cell.paragraphs[para_index]
    else:
        # Target paragraph doesn't exist — add a new one
        # This preserves the label in para 0
        para = cell.add_paragraph()

    # If the paragraph already has runs, set the first one
    if para.runs:
        para.runs[0].text = value
        for r in para.runs[1:]:
            r.text = ""
    else:
        # Add a new run matching the original font style
        run = para.add_run(value)
        run.font.name = font_name
        run.font.size = font_size  # EMU: 127000 = 10pt


def write_amount_to_cell(cell, value):
    """Write a dollar amount into an empty fee cell (column 1 of Table 6)."""
    value = value or ""
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = value
        for r in para.runs[1:]:
            r.text = ""
    else:
        run = para.add_run(value)
        run.font.name = "Times New Roman"
        run.font.size = 127000


def build_address(contact):
    parts = [
        contact.get("address"),
        contact.get("city"),
        contact.get("state"),
        contact.get("zip"),
        contact.get("country"),
    ]
    return ", ".join(p for p in parts if p)


def fill_student_info(table, contact):
    """Fill Table 0 — Student Information (8 rows x 3 cols).

    Each cell has a label in Para 0 and an empty Para 1 for the value.
    Cells with dropdowns (Immigration Status, Gender, Study Permit) are
    handled by fill_all_dropdowns() — skip them here.
    """
    # Row 2: Last Name | Usual First Name | First Name & Middle Name
    write_value_to_cell(table.rows[2].cells[0], contact.get("lastname") or "")
    write_value_to_cell(table.rows[2].cells[1], contact.get("firstname") or "")
    write_value_to_cell(table.rows[2].cells[2], contact.get("firstname") or "")

    # Row 3: Email | Phone | PEN
    write_value_to_cell(table.rows[3].cells[0], contact.get("email") or "")
    write_value_to_cell(table.rows[3].cells[1], contact.get("phone") or contact.get("mobilephone") or "")
    # PEN [3,2] — only 1 para, advisor fills later

    # Row 4: Mailing Address (spans 2 cells) | Date of Birth
    address = build_address(contact)
    write_value_to_cell(table.rows[4].cells[0], address)
    write_value_to_cell(table.rows[4].cells[2], contact.get("date_of_birth") or "")

    # Row 5: Mailing Address in Canada (spans full width) — leave blank

    # Row 6: Immigration Status [6,0] | Citizenship [6,1] | Gender [6,2]
    # Immigration Status and Gender have DROPDOWNS — handled by fill_all_dropdowns()
    # Only fill Citizenship here (has empty Para 1)
    write_value_to_cell(table.rows[6].cells[1], contact.get("citizenship") or "")

    # Row 7: Study permit — has DROPDOWN, handled by fill_all_dropdowns()


def fill_program_info(table, program, intake, delivery_method=""):
    """Fill Table 2 — Program Information (6 rows x 5 cols).

    Row 1: Program Title | Hours | Weeks | Start Date | End Date
    Row 2: Credential (DROPDOWN) | Schedule (DROPDOWN) | ...
    Cells with dropdowns are handled by fill_all_dropdowns().

    Hours, weeks, and delivery method come from the INTAKE, not the program.
    """
    # Hours and weeks from intake (v2 schema)
    hours = intake.get("hours", "") if intake else ""
    weeks = intake.get("weeks", "") if intake else ""

    # Row 1: all have empty Para 1 except Hours [1,1] which has only 1 para
    write_value_to_cell(table.rows[1].cells[0], program.get("program_name") or "")
    write_value_to_cell(table.rows[1].cells[1], str(hours))
    write_value_to_cell(table.rows[1].cells[2], str(weeks))
    write_value_to_cell(table.rows[1].cells[3], intake.get("intake_date", "") if intake else "")
    write_value_to_cell(table.rows[1].cells[4], intake.get("end_date", "") if intake else "")

    # Row 2: Credential [2,0] and Schedule [2,1] have DROPDOWNS
    # Handled by fill_all_dropdowns() — don't write here


def fill_fees_table(table, fee_items, amount_col=None):
    """Fill Table 6 — Program Costs (7 rows x 2 cols).

    Deep-copies row 1 (first data row) for each fee to preserve all XML
    formatting (borders, shading, font, cell widths). Only changes text.

    fee_items: list of {name, amount, isTuition} dicts from resolve_fee_amounts()
    amount_col: ignored in v3 schema (kept for signature compatibility)
    """
    tbl = table._tbl
    num_rows = len(table.rows)

    # Save XML templates before modifying
    header_tr = table.rows[0]._tr            # "PROGRAM COSTS" header
    data_template_tr = copy.deepcopy(table.rows[1]._tr)  # Data row format
    total_template_tr = copy.deepcopy(table.rows[num_rows - 1]._tr)  # Total row

    # Remove all rows except header (row 0)
    rows_to_remove = [table.rows[i]._tr for i in range(1, num_rows)]
    for tr in rows_to_remove:
        tbl.remove(tr)

    # Add one row per fee item (deep-copied from data row template)
    grand_total = 0.0
    for fee in fee_items:
        amount = float(fee.get("amount", 0) or 0)
        grand_total += amount
        amount_str = f"${amount:,.2f}"

        new_tr = copy.deepcopy(data_template_tr)
        tbl.append(new_tr)
        row = table.rows[-1]

        # Cell 0: fee name — replace existing text in runs
        fee_name = fee.get("name", "")
        cell0 = row.cells[0]
        if cell0.paragraphs[0].runs:
            cell0.paragraphs[0].runs[0].text = fee_name
            for r in cell0.paragraphs[0].runs[1:]:
                r.text = ""
        else:
            run = cell0.paragraphs[0].add_run(fee_name)
            run.font.name = "Times New Roman"
            run.font.size = 127000

        # Cell 1: amount
        write_amount_to_cell(row.cells[1], amount_str)

    # Add total row back
    new_total_tr = copy.deepcopy(total_template_tr)
    tbl.append(new_total_tr)
    total_row = table.rows[-1]

    # Set total amount in cell 1
    write_amount_to_cell(total_row.cells[1], f"${grand_total:,.2f}")


def fill_docx(template_path, output_path, contact, program, fees, amount_col, intake, delivery_method=""):
    """Fill the DOCX by writing values into empty cells/paragraphs.

    Only adds text to designated empty spots. Never modifies labels,
    borders, colors, fonts, images, headers, footers, or section layout.
    """
    doc = Document(str(template_path))

    # Table 0: Student Information
    if len(doc.tables) > 0:
        fill_student_info(doc.tables[0], contact)

    # Strip all dropdown controls and replace with plain text values
    replace_all_dropdowns(doc, contact, program, intake)

    # Table 2: Program Information
    if len(doc.tables) > 2:
        fill_program_info(doc.tables[2], program, intake, delivery_method=delivery_method)

    # Remove sections that are no longer in the contract body:
    # - Table 1: Voluntary Disclosure (now in PandaDoc signing template)
    # - Tables 3-5: Program Outline placeholder, Work Experience, Regulatory Requirement
    #   (replaced by the actual program outline PDF added as a PandaDoc section)
    # - Table 19: Signature table (handled by PandaDoc signing template)
    # Remove in REVERSE order so indices don't shift
    tables_to_remove = []
    for idx in [19, 5, 4, 3, 1]:
        if idx < len(doc.tables):
            tables_to_remove.append(doc.tables[idx]._tbl)
    for tbl_element in tables_to_remove:
        parent = tbl_element.getparent()
        if parent is not None:
            parent.remove(tbl_element)

    # Table 6: Program Costs (index shifted after removals, find by content)
    # After removing tables 1,3,4,5 the fees table moved. Find it by looking
    # for the table whose first cell contains "PROGRAM COSTS" or similar.
    fees_table = None
    for t in doc.tables:
        try:
            first_cell_text = t.rows[0].cells[0].text.strip().upper()
            if "PROGRAM COSTS" in first_cell_text or "COST" in first_cell_text:
                fees_table = t
                break
        except (IndexError, AttributeError):
            continue
    if fees_table:
        fill_fees_table(fees_table, fees, amount_col)

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def split_docx_at_fees(docx_path):
    """Split a filled DOCX into two parts at the PROGRAM COSTS table.

    Returns (part1_path, part2_path) where:
    - part1: Student Info + Program Info (everything before fees)
    - part2: Fees + Payment Terms + Refund + T&Cs (everything from fees onward)

    If the fees table is not found, returns (docx_path, None) — no split.
    """
    from docx.oxml.ns import qn as docx_qn

    # Part 1: everything before the fees table
    doc1 = Document(str(docx_path))
    body1 = doc1.element.body

    fees_tbl = None
    for tbl in doc1.tables:
        try:
            text = tbl.rows[0].cells[0].text.strip().upper()
            if "PROGRAM COSTS" in text or "COST" in text:
                fees_tbl = tbl._tbl
                break
        except (IndexError, AttributeError):
            continue

    if fees_tbl is None:
        return docx_path, None

    # Remove everything from fees table onward
    found = False
    to_remove = []
    for child in list(body1):
        if child is fees_tbl:
            found = True
        if found:
            to_remove.append(child)
    for elem in to_remove:
        body1.remove(elem)

    part1_path = docx_path.parent / (docx_path.stem + "_part1.docx")
    doc1.save(str(part1_path))

    # Part 2: fees table and everything after
    doc2 = Document(str(docx_path))
    body2 = doc2.element.body

    fees_tbl2 = None
    for tbl in doc2.tables:
        try:
            text = tbl.rows[0].cells[0].text.strip().upper()
            if "PROGRAM COSTS" in text or "COST" in text:
                fees_tbl2 = tbl._tbl
                break
        except (IndexError, AttributeError):
            continue

    if fees_tbl2 is not None:
        to_remove = []
        for child in list(body2):
            if child is fees_tbl2:
                break
            to_remove.append(child)
        for elem in to_remove:
            body2.remove(elem)

    part2_path = docx_path.parent / (docx_path.stem + "_part2.docx")
    doc2.save(str(part2_path))

    return part1_path, part2_path


# ---------------------------------------------------------------------------
# PandaDoc — Create from template + add sections
# ---------------------------------------------------------------------------

SIGNING_TEMPLATE_ID = os.environ.get("PANDADOC_SIGNING_TEMPLATE_ID", "DsoTNroxu5qX332aioZVjK")
OUTLINE_MAP_PATH = DATA_DIR / "program-outline-map.csv"
OUTLINE_CACHE_DIR = Path(os.environ.get("OUTLINE_CACHE_DIR", Path(tempfile.gettempdir()) / "wcc_outlines"))


def load_outline_map():
    """Load program -> {filename, google_drive_file_id} mapping from CSV."""
    mapping = {}
    if not OUTLINE_MAP_PATH.exists():
        print(f"  Warning: outline map not found at {OUTLINE_MAP_PATH}, skipping.")
        return mapping
    with open(OUTLINE_MAP_PATH, newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            program_key = row["program_name"].strip().lower()
            drive_id = row.get("google_drive_file_id", "").strip()
            filename = row.get("outline_filename", "").strip()
            if drive_id:
                mapping[program_key] = {
                    "filename": filename,
                    "google_drive_file_id": drive_id,
                }
    return mapping


def download_outline_from_drive(file_id):
    """Download a program outline PDF from Google Drive.

    Uses the public export URL for link-shared files (no API key needed).
    Caches downloaded files in /tmp/outlines/{file_id}.pdf to avoid re-downloading.
    Returns the local file path on success, or None on failure.
    """
    if not re.match(r'^[A-Za-z0-9_-]{10,60}$', file_id):
        print(f"  Invalid file_id format: {file_id}")
        return None

    # Clean up cached files older than 24 hours
    if OUTLINE_CACHE_DIR.exists():
        now = time.time()
        for cached_file in OUTLINE_CACHE_DIR.glob("*.pdf"):
            try:
                if now - cached_file.stat().st_mtime > 86400:  # 24 hours
                    cached_file.unlink()
            except Exception:
                pass

    OUTLINE_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    cached_path = OUTLINE_CACHE_DIR / f"{file_id}.pdf"

    if cached_path.exists():
        return cached_path

    url = f"https://drive.google.com/uc?export=download&id={file_id}"
    req = urllib.request.Request(url)
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = resp.read()
        with open(cached_path, "wb") as f:
            f.write(data)
        return cached_path
    except Exception as e:
        print(f"  Warning: Failed to download outline from Drive (file_id={file_id}): {e}")
        return None


def get_outline_path(outline_map, program_name):
    """Look up and download the PDF outline for a program from Google Drive.

    Returns the local cached file path, or None if unavailable.
    """
    entry = outline_map.get(program_name.lower())
    if not entry:
        return None

    file_id = entry.get("google_drive_file_id")
    if not file_id:
        return None

    return download_outline_from_drive(file_id)


def create_from_signing_template(doc_name, advisor, student_email, student_first, student_last):
    """Step 1: Create a document from the signing template.

    The template has pre-placed signature + date fields for Advisor and Client roles.
    Signing order: Advisor (1) first, Client/Student (2) second.
    """
    payload = {
        "name": doc_name,
        "template_uuid": SIGNING_TEMPLATE_ID,
        "recipients": [
            {
                "email": advisor.get("email", ""),
                "first_name": advisor.get("firstName", ""),
                "last_name": advisor.get("lastName", ""),
                "role": "Advisor",
                "signing_order": 1,
            },
            {
                "email": student_email,
                "first_name": student_first,
                "last_name": student_last,
                "role": "Client",
                "signing_order": 2,
            },
        ],
    }

    return api_post(
        f"{PANDADOC_BASE_URL}/documents",
        payload,
        PANDADOC_API_KEY,
        token_type="API-Key",
    )


def wait_for_document_draft(doc_id, max_wait=30):
    """Wait for document to move from 'uploaded' to 'draft' status."""
    for _ in range(max_wait):
        status = api_get(
            f"{PANDADOC_BASE_URL}/documents/{doc_id}",
            PANDADOC_API_KEY,
            token_type="API-Key",
        )
        doc_status = status.get("status", "")
        if doc_status == PANDADOC_STATUS_DRAFT:
            return True
        if PANDADOC_STATUS_ERROR_PREFIX in doc_status:
            raise Exception(f"Document error: {doc_status}")
        time.sleep(1)
    raise Exception(f"Document did not reach draft status within {max_wait}s")


def add_section_from_file(doc_id, file_path, section_name="Section", mime_type=None):
    """Step 2/3: Add a file as a new section to an existing document.

    POST /public/v1/documents/{id}/sections/uploads
    """
    if mime_type is None:
        if str(file_path).endswith(".pdf"):
            mime_type = "application/pdf"
        else:
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    with open(file_path, "rb") as f:
        file_data = f.read()

    return api_post_multipart(
        f"{PANDADOC_BASE_URL}/documents/{doc_id}/sections/uploads",
        fields={"data": json.dumps({"name": section_name})},
        files={"file": (Path(file_path).name, file_data, mime_type)},
        token=PANDADOC_API_KEY,
        token_type="API-Key",
    )


def link_to_hubspot(doc_id, deal_id):
    """Link the PandaDoc document to the HubSpot deal."""
    try:
        api_post(
            f"{PANDADOC_BASE_URL}/documents/{doc_id}/linked-objects",
            {"provider": "hubspot", "entity_id": deal_id, "entity_type": "deal"},
            PANDADOC_API_KEY,
            token_type="API-Key",
        )
        return True
    except Exception as e:
        print(f"  Warning: Could not link to HubSpot deal: {e}")
        return False


# ---------------------------------------------------------------------------
# Template validation
# ---------------------------------------------------------------------------

def validate_template(template_path):
    """Validate the DOCX template has expected table structure."""
    doc = Document(str(template_path))
    tables = doc.tables

    # Check minimum table count
    if len(tables) < 6:
        print(f"WARNING: Template has only {len(tables)} tables (expected 6+)")
        return False

    # Check for PROGRAM COSTS table
    has_costs = False
    for t in tables:
        try:
            text = t.rows[0].cells[0].text.strip().upper()
            if "PROGRAM COSTS" in text or "COST" in text:
                has_costs = True
                break
        except (IndexError, AttributeError):
            continue

    if not has_costs:
        print("WARNING: No 'PROGRAM COSTS' table found in template. DOCX split will not work.")
        return False

    # Check for Student Information table
    has_student = False
    try:
        if "STUDENT" in tables[0].rows[0].cells[0].text.strip().upper():
            has_student = True
    except (IndexError, AttributeError):
        pass

    if not has_student:
        print("WARNING: Table 0 does not appear to be Student Information.")

    return True


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) < 2:
        print("Usage: python generate-contract-v2.py <hubspot-deal-id> [intake-date]")
        print("")
        print("Example:")
        print("  python generate-contract-v2.py 58197228338")
        print("  python generate-contract-v2.py 58197228338 2026-04-14")
        sys.exit(1)

    deal_id = sys.argv[1]
    selected_intake_date = sys.argv[2] if len(sys.argv) > 2 and not sys.argv[2].startswith("--") else ""

    # Check for --data-file argument
    data_file_path = None
    for i, arg in enumerate(sys.argv):
        if arg == "--data-file" and i + 1 < len(sys.argv):
            data_file_path = sys.argv[i + 1]
            break

    print("=" * 50)
    print("WCC Contract Generator v2")
    print("=" * 50)

    # Validate template structure
    if not TEMPLATE_PATH.exists():
        print(f"ERROR: Template not found: {TEMPLATE_PATH}")
        sys.exit(1)
    validate_template(TEMPLATE_PATH)

    # Step 1: Load program/fee data
    print("\n1. Loading program data...")
    if data_file_path and os.path.exists(data_file_path):
        # Use pre-cached data from backend (avoids Google Sheets API calls)
        with open(data_file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        print(f"   {len(data.get('programs', []))} programs, {len(data.get('fees', []))} fee items (from cache)")
    else:
        # Fallback: load from Google Sheets or CSV directly
        data = load_data()
        print(f"   {len(data.get('programs', []))} programs, {len(data.get('fees', []))} fee items")

    # Step 2: Get HubSpot deal + contact
    print(f"\n2. Fetching HubSpot deal {deal_id}...")
    hs = get_deal_data(deal_id)
    contact = hs["contact"]
    deal = hs["deal"]

    program_name = deal.get("program_of_study") or contact.get("program_of_study", "")
    residence_status = contact.get("residence_status", "")
    domestic = is_domestic(residence_status)

    if domestic is None:
        print(f"\nERROR: residence_status is not set on the contact.")
        print(f"   Set it in HubSpot first (Canadian Citizen, Permanent Resident, International Study Permit, etc.)")
        print(f"   Contact: [redacted] (deal {deal_id})")
        sys.exit(1)

    tier = "Domestic" if domestic else "International"

    print(f"   Student: {contact.get('firstname', '')[0]}*** {contact.get('lastname', '')[0]}***")
    print(f"   Email:   [redacted]")
    print(f"   Program: {program_name}")
    print(f"   Status:  {residence_status} -> {tier}")

    if not program_name:
        print("\nERROR: No program_of_study on deal or contact.")
        sys.exit(1)

    # Step 3: Look up program (identity only in v2 schema)
    print(f"\n3. Looking up program data...")
    program = lookup_program(data["programs"], program_name)
    if not program:
        print(f"   ERROR: Program '{program_name}' not found in data.")
        sys.exit(1)
    print(f"   Code: {program['program_code']} | {program['credential']}")

    # Step 3b: Find the intake
    intake = None
    if selected_intake_date:
        # Find specific intake by date
        for i in data.get("intakes", []):
            if (i.get("program_name", "").lower() == program_name.lower()
                    and i.get("intake_date") == selected_intake_date):
                intake = i
                break
        if intake:
            print(f"   Selected intake: {intake['intake_date']} -> {intake['end_date']} at {intake.get('campus', '?')}")
        else:
            print(f"   Warning: Intake {selected_intake_date} not found, using date for fee matching only")
    else:
        # Fallback: pick next open intake
        intake = lookup_next_intake(data.get("intakes", []), program_name)
        if intake:
            print(f"   Next intake: {intake['intake_date']} -> {intake['end_date']} at {intake.get('campus', '?')}")
        else:
            print("   No upcoming intake (advisor fills dates in PandaDoc)")

    # Determine delivery method based on residency
    delivery_method = ""
    if intake:
        if domestic:
            delivery_method = intake.get("domestic_delivery_method", "")
        else:
            delivery_method = intake.get("international_delivery_method", "")
        if delivery_method:
            print(f"   Delivery: {delivery_method} ({tier})")

    # Step 3c: Match fees using effective_from logic (column-based schema v3)
    reference_date = selected_intake_date or (intake.get("intake_date") if intake else "") or datetime.now(PST).strftime("%Y-%m-%d")
    residency_tier = "domestic" if domestic else "international"
    matched_fees, effective_from = match_fees(data["fees"], program_name, reference_date, residency_tier)
    if not matched_fees:
        print(f"   ERROR: No fees found for '{program_name}' / {residency_tier} effective on {reference_date}.")
        sys.exit(1)

    fee_items, total = resolve_fee_amounts(matched_fees, domestic)

    print(f"   {len(fee_items)} fee items ({tier}, effective from {effective_from}):")
    for fi in fee_items:
        print(f"     {fi['name']}: ${fi['amount']:,.2f}")
    print(f"     TOTAL: ${total:,.2f}")

    # Step 4: Fill the DOCX
    print(f"\n4. Filling DOCX template...")
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    safe_name = re.sub(r'[^\w\s-]', '', f"{contact.get('lastname', 'Unknown')}_{program.get('program_code', 'UNK')}")
    output_filename = f"{safe_name}_{tier}_{timestamp}.docx"
    output_path = OUTPUT_DIR / output_filename

    fill_docx(
        template_path=TEMPLATE_PATH,
        output_path=output_path,
        contact=contact,
        program=program,
        fees=fee_items,
        amount_col=None,
        intake=intake,
        delivery_method=delivery_method,
    )
    print(f"   Saved: {output_path}")

    # Step 5: Look up outline + split DOCX if outline exists
    advisor = hs["advisor"]
    print(f"   Advisor: {advisor.get('firstName', '')} {advisor.get('lastName', '')}")

    # Use cached outline map from backend if available, otherwise load from CSV
    if data_file_path and "outline_map" in data:
        outline_map = data["outline_map"]
    else:
        outline_map = load_outline_map()
    outline_path = get_outline_path(outline_map, program_name)

    doc_name = f"{contact.get('firstname', '')} {contact.get('lastname', '')} - {program_name} Enrollment Contract"

    recipients = [
        {
            "email": advisor.get("email", ""),
            "first_name": advisor.get("firstName", ""),
            "last_name": advisor.get("lastName", ""),
            "role": "Advisor",
            "signing_order": 1,
        },
        {
            "email": contact.get("email", ""),
            "first_name": contact.get("firstname", ""),
            "last_name": contact.get("lastname", ""),
            "role": "Client",
            "signing_order": 2,
        },
    ]

    if outline_path:
        # Split DOCX: Part 1 (Student Info + Program Info) → Outline PDF → Part 2 (Fees + rest)
        print(f"\n5. Splitting DOCX for outline insertion...")
        part1_path, part2_path = split_docx_at_fees(output_path)

        if part2_path:
            print(f"   Part 1: {part1_path.name}")
            print(f"   Outline: {outline_path.name}")
            print(f"   Part 2: {part2_path.name}")

            # Upload Part 1 as main document
            print(f"\n6. Uploading Part 1 (Student Info + Program Info)...")
            with open(part1_path, "rb") as f:
                file_data = f.read()
            metadata = {"name": doc_name, "recipients": recipients, "parse_form_fields": False}
            result = api_post_multipart(
                f"{PANDADOC_BASE_URL}/documents",
                fields={"data": json.dumps(metadata)},
                files={"file": (part1_path.name, file_data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                token=PANDADOC_API_KEY, token_type="API-Key",
            )
            doc_id = result.get("id", "")
            print(f"   Document ID: {doc_id}")

            print(f"\n   Waiting for draft status...")
            wait_for_document_draft(doc_id)

            # Add outline PDF as section
            print(f"\n7. Adding program outline PDF...")
            add_section_from_file(doc_id, outline_path, section_name=f"Program Outline - {program_name}")
            print(f"   Outline added: {outline_path.name}")

            # Add Part 2 as section (Fees + Payment Terms + etc.)
            print(f"\n   Adding Part 2 (Fees + Terms)...")
            add_section_from_file(doc_id, part2_path, section_name="Program Costs and Terms")
            print(f"   Part 2 added")
        else:
            # Split failed — upload full DOCX and add outline after
            print(f"   Split failed, uploading full DOCX...")
            with open(output_path, "rb") as f:
                file_data = f.read()
            metadata = {"name": doc_name, "recipients": recipients, "parse_form_fields": False}
            result = api_post_multipart(
                f"{PANDADOC_BASE_URL}/documents",
                fields={"data": json.dumps(metadata)},
                files={"file": (output_path.name, file_data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                token=PANDADOC_API_KEY, token_type="API-Key",
            )
            doc_id = result.get("id", "")
            wait_for_document_draft(doc_id)
            add_section_from_file(doc_id, outline_path, section_name=f"Program Outline - {program_name}")
    else:
        # No outline — upload full DOCX as-is
        print(f"\n5. No outline found for '{program_name}', uploading full DOCX...")
        with open(output_path, "rb") as f:
            file_data = f.read()
        metadata = {"name": doc_name, "recipients": recipients, "parse_form_fields": False}
        result = api_post_multipart(
            f"{PANDADOC_BASE_URL}/documents",
            fields={"data": json.dumps(metadata)},
            files={"file": (output_path.name, file_data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            token=PANDADOC_API_KEY, token_type="API-Key",
        )
        doc_id = result.get("id", "")
        print(f"   Document ID: {doc_id}")
        print(f"\n6. Waiting for draft status...")
        wait_for_document_draft(doc_id)

    # Step 8: Add signing template as last section (with role mapping)
    print(f"\n7. Adding signing page as section...")
    section_payload = {
        "template_uuid": SIGNING_TEMPLATE_ID,
        "name": "Signatures",
        "recipients": [
            {
                "email": advisor.get("email", ""),
                "first_name": advisor.get("firstName", ""),
                "last_name": advisor.get("lastName", ""),
                "role": "Advisor",
                "signing_order": 1,
            },
            {
                "email": contact.get("email", ""),
                "first_name": contact.get("firstname", ""),
                "last_name": contact.get("lastname", ""),
                "role": "Client",
                "signing_order": 2,
            },
        ],
    }
    signing_result = api_post(
        f"{PANDADOC_BASE_URL}/documents/{doc_id}/sections/uploads",
        section_payload,
        PANDADOC_API_KEY,
        token_type="API-Key",
    )
    print(f"   Signing page added: {signing_result.get('status', '?')}")

    # Step 9: Link to HubSpot deal
    print(f"\n9. Linking to HubSpot deal...")
    link_to_hubspot(doc_id, deal_id)

    # Contract logging is handled by the backend server (backend-server.py)

    doc_url = f"https://app.pandadoc.com/a/#/documents/{doc_id}"

    # Cleanup: remove local DOCX files containing student PII
    try:
        if output_path.exists():
            output_path.unlink()
        # Also clean up split parts if they exist
        part1 = output_path.parent / (output_path.stem + "_part1.docx")
        part2 = output_path.parent / (output_path.stem + "_part2.docx")
        if part1.exists():
            part1.unlink()
        if part2.exists():
            part2.unlink()
    except Exception as e:
        print(f"   Warning: Could not clean up local files: {e}")

    # Done
    print("\n" + "=" * 50)
    print("CONTRACT CREATED SUCCESSFULLY")
    print("=" * 50)
    print(f"Document ID: {doc_id}")
    print(f"View:     {doc_url}")
    print(f"Fee tier: {tier}")
    if effective_from:
        print(f"Fees effective from: {effective_from}")
    if delivery_method:
        print(f"Delivery: {delivery_method}")
    print("")
    print("Next steps:")
    print("  1. Open in PandaDoc")
    print("  2. Review all data")
    print("  3. Sign as Advisor")
    print("  4. Click Send -> Student receives for signature")


if __name__ == "__main__":
    main()
